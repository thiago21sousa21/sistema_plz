import os
import shutil
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

arquivo_atual = os.path.abspath(__file__)
diretorio_atual = os.path.dirname(arquivo_atual)
caminho_imagens = fr"C:\Users\usimp\OneDrive\Imagens"
caminho_laudo_modelo = os.path.join(diretorio_atual, '..', 'modelo_laudo', 'laudo_modelo.docx')
caminho_copia = os.path.join(diretorio_atual, '..', 'laudo_criado.docx')
caminho_cameras = os.path.join(diretorio_atual, '..', 'images', 'mapas_pontos_vd')

def listar_diretorio(caminho):
    diretorio = os.listdir(caminho)
    return diretorio

def pegar_fotos(q=5):
    lista_fotos = listar_diretorio(caminho_imagens)[-1*q:]
    return(lista_fotos)

def criar_copia_laudo_modelo(caminho_origem, caminho_destino):
    shutil.copy(caminho_origem, caminho_destino)

def deletar_ultima_copia_laudo(caminho):
    try:
        os.remove(caminho)
        print("Arquivo deletado com sucesso!")
    except FileNotFoundError:
        print("O arquivo não existe.")
    except PermissionError:
        print("Permissão negada para deletar o arquivo.")
    except Exception as e:
        print(f"Ocorreu um erro: {e}")

def inserir_local():
    opc = input('Gostaria de inserir a camera? [s/n]: ').strip()[0].lower()
    while opc not in 'ns':
        opc = input('apenas n ou s: ').strip()[0].lower()
    if opc == 's':
        numero_camera = input('Digite o número da camera: ')
        return numero_camera
    return False


def converter_docx_para_pdf(sigla_fiscal, numero_laudo):
    from docx2pdf import convert
    """
    Procura por um arquivo .docx e o converte para .pdf com um novo nome fornecido pelo usuário.
    """
    #nome_docx = input("Digite o nome do arquivo .docx (sem a extensão): ").strip()
    caminho_docx =  os.path.join(diretorio_atual, "..", "laudo_criado.docx") 

    if not os.path.exists(caminho_docx):
        print(f"Arquivo '{caminho_docx}' não encontrado!")
        return

    #numeor_auto = input("Qual numero quer dar laudo? (sem extensão): ").strip().zfill(3)
    caminho_pdf =  os.path.join(diretorio_atual, "..", f"LF{sigla_fiscal[2:]}-{str(numero_laudo).zfill(3)}.pdf") 

    try:
        convert(caminho_docx, caminho_pdf)
        print(f"Arquivo convertido com sucesso: {caminho_pdf}")
    except Exception as e:
        print("Erro ao converter o arquivo:", e)

def inserir_imagens():
    deletar_ultima_copia_laudo( caminho_copia)
    criar_copia_laudo_modelo(caminho_laudo_modelo, caminho_copia)
    doc = Document(caminho_copia)
##NESSE LUGAR TENHO QUE PEDIR O NOME 

    if doc.paragraphs:
        fiscal = input("Digite a sigla do fiscal: ").upper()
        numero_laudo = input("Digite o número do laudo: ").upper().zfill(3)
        nome_autuado = input("Digite o nome do autuado: ").upper()

        nova_frase = f'{fiscal} - {numero_laudo}/2025 - {nome_autuado}'
        primeiro_paragrafo = doc.paragraphs[0]
        primeiro_paragrafo.clear()
        novo_primeiro_paragrafo = primeiro_paragrafo.add_run(nova_frase)
        novo_primeiro_paragrafo.font.name = 'calibri'
        novo_primeiro_paragrafo.font.size = Pt(11)
        novo_primeiro_paragrafo.bold = True



    fotos = pegar_fotos(int(input("Digite quantas fotos quer pegar: ")))
    print(fotos)
    for f in fotos:
        foto = fr'{caminho_imagens}\{f}'
        img_paragraph = doc.add_paragraph()
        img_run = img_paragraph.add_run()
        img_run.add_picture(foto, width=Cm(15), height=Cm(7))
        img_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    numero_camera = inserir_local()
    if numero_camera:
        caminho_camera = os.path.join(caminho_cameras, f'cam{numero_camera}.png')
        print(caminho_camera)
        img_paragraph = doc.add_paragraph()
        img_run = img_paragraph.add_run()
        img_run.add_picture(caminho_camera,  width=Cm(15), height=Cm(8))
        img_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(caminho_copia)
    converter_docx_para_pdf(fiscal, numero_laudo)



